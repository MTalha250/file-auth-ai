import os
import fitz  # PyMuPDF
import pytesseract
from PIL import Image
import cv2
import numpy as np
from pathlib import Path
from typing import Dict, Optional, List, Tuple
import io

# --- DOCX Text Extractor ---
class DocxTextExtractor:
    """
    DOCX text extraction pipeline that extracts text and images with OCR.
    Similar structure to PDFTextExtractor.
    """
    
    def __init__(self, tesseract_path: Optional[str] = None, output_dir: str = "extracted_texts"):
        """Initialize the DOCX text extractor."""
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(exist_ok=True)
        
        # Configure tesseract path if provided
        if tesseract_path:
            pytesseract.pytesseract.tesseract_cmd = tesseract_path
        
        # Setup logging
        import logging
        logging.basicConfig(
            level=logging.INFO,
            format='%(asctime)s - %(levelname)s - %(message)s'
        )
        self.logger = logging.getLogger(__name__)
    
    def diagnose_docx(self, docx_path: str) -> Dict[str, any]:
        """Diagnose DOCX structure to understand content types."""
        try:
            from docx import Document
            import zipfile
            
            doc = Document(docx_path)
            diagnosis = {
                'total_paragraphs': len(doc.paragraphs),
                'total_tables': len(doc.tables),
                'images_found': 0,
                'sections': len(doc.sections)
            }
            
            # Count images by checking the ZIP structure
            with zipfile.ZipFile(docx_path, 'r') as docx_zip:
                image_files = [f for f in docx_zip.namelist() 
                             if f.startswith('word/media/') and 
                             f.lower().endswith(('.png', '.jpg', '.jpeg', '.gif', '.bmp', '.tiff'))]
                diagnosis['images_found'] = len(image_files)
                diagnosis['image_files'] = image_files
            
            self.logger.info(f"DOCX diagnosis: {diagnosis['total_paragraphs']} paragraphs, "
                           f"{diagnosis['total_tables']} tables, {diagnosis['images_found']} images")
            
            return diagnosis
            
        except Exception as e:
            self.logger.error(f"DOCX diagnosis failed: {e}")
            return {'error': str(e)}
    
    def extract_native_text(self, docx_path: str) -> Dict:
        """Extract native text from DOCX using python-docx."""
        try:
            from docx import Document
            doc = Document(docx_path)
            
            # Extract paragraphs
            paragraphs = []
            for i, paragraph in enumerate(doc.paragraphs):
                if paragraph.text.strip():
                    paragraphs.append({
                        'paragraph_num': i + 1,
                        'text': paragraph.text.strip()
                    })
            
            # Extract tables
            tables = []
            for table_idx, table in enumerate(doc.tables):
                table_data = []
                for row_idx, row in enumerate(table.rows):
                    row_data = []
                    for cell in row.cells:
                        row_data.append(cell.text.strip())
                    table_data.append(row_data)
                
                tables.append({
                    'table_num': table_idx + 1,
                    'data': table_data,
                    'rows': len(table_data),
                    'cols': len(table_data[0]) if table_data else 0
                })
            
            # Extract headers and footers
            headers = []
            footers = []
            for section_idx, section in enumerate(doc.sections):
                if section.header:
                    for para in section.header.paragraphs:
                        if para.text.strip():
                            headers.append(para.text.strip())
                if section.footer:
                    for para in section.footer.paragraphs:
                        if para.text.strip():
                            footers.append(para.text.strip())
            
            text_data = {
                'paragraphs': paragraphs,
                'tables': tables,
                'headers': headers,
                'footers': footers,
                'total_paragraphs': len(paragraphs),
                'total_tables': len(tables)
            }
            
            self.logger.info(f"Extracted {len(paragraphs)} paragraphs and {len(tables)} tables from DOCX")
            return text_data
            
        except ImportError:
            self.logger.error("python-docx not installed. Install with: pip install python-docx")
            return {'error': 'python-docx not installed'}
        except Exception as e:
            self.logger.error(f"Failed to extract native text: {e}")
            # Fallback to docx2txt if available
            try:
                import docx2txt
                text = docx2txt.process(docx_path)
                return {
                    'paragraphs': [{'paragraph_num': 1, 'text': text}],
                    'tables': [],
                    'headers': [],
                    'footers': [],
                    'fallback_used': True
                }
            except:
                return {'error': str(e)}
    
    def extract_images_from_docx(self, docx_path: str) -> List[Dict]:
        """Extract images from DOCX file and save them to extracted_images folder."""
        images = []
        
        # Create extracted_images folder
        extracted_images_dir = self.output_dir / "extracted_images"
        extracted_images_dir.mkdir(exist_ok=True)
        
        try:
            import zipfile
            
            # DOCX is essentially a ZIP file
            with zipfile.ZipFile(docx_path, 'r') as docx_zip:
                # Find all image files in the media folder
                image_files = [f for f in docx_zip.namelist() 
                             if f.startswith('word/media/') and 
                             f.lower().endswith(('.png', '.jpg', '.jpeg', '.gif', '.bmp', '.tiff'))]
                
                self.logger.info(f"Found {len(image_files)} images in DOCX")
                
                for img_index, img_file in enumerate(image_files):
                    try:
                        # Extract image data
                        image_bytes = docx_zip.read(img_file)
                        
                        # Filter out tiny images (likely decorative elements)
                        if len(image_bytes) < 1000:
                            self.logger.debug(f"Skipping tiny image {img_index + 1} (size: {len(image_bytes)} bytes)")
                            continue
                        
                        # Convert to PIL Image
                        image = Image.open(io.BytesIO(image_bytes))
                        
                        # Filter out very small images by dimensions
                        if image.width < 50 or image.height < 50:
                            self.logger.debug(f"Skipping small image {img_index + 1} (dimensions: {image.width}x{image.height})")
                            continue
                        
                        # Save the image to disk
                        original_name = Path(img_file).name
                        img_filename = f"docx_img{img_index + 1}_{original_name}"
                        img_path = extracted_images_dir / img_filename
                        image.save(img_path)
                        
                        image_data = {
                            "image": image,
                            "index": img_index + 1,
                            "size": (image.width, image.height),
                            "type": "image",
                            "saved_path": str(img_path),
                            "filename": img_filename,
                            "original_name": original_name
                        }
                        
                        images.append(image_data)
                        self.logger.info(f"Successfully extracted and saved image {img_index + 1} (size: {image.width}x{image.height}, {len(image_bytes)} bytes) -> {img_filename}")
                        
                    except Exception as e:
                        self.logger.warning(f"Failed to extract image {img_index + 1} from DOCX: {e}")
            
            self.logger.info(f"Successfully extracted {len(images)} valid images from DOCX")
            return images
            
        except Exception as e:
            self.logger.error(f"Failed to extract images from DOCX: {e}")
            return []
    
    def extract_text_from_image(self, image: Image.Image) -> str:
        """
        Extract text from image using Tesseract OCR - reuse the same method as PDFTextExtractor.
        """
        try:
            # Test if Tesseract is available
            try:
                version = pytesseract.get_tesseract_version()
                self.logger.debug(f"Using Tesseract version: {version}")
            except Exception as e:
                self.logger.error(f"Tesseract not available: {e}")
                return ""
            
            all_results = []
            
            # Try different OCR configurations
            configs = [
                ('default', r'--oem 3 --psm 6'),
                ('psm7', r'--oem 3 --psm 7'),
                ('psm8', r'--oem 3 --psm 8'),
                ('psm11', r'--oem 3 --psm 11'),
                ('psm13', r'--oem 3 --psm 13')
            ]
            
            for name, config in configs:
                try:
                    text = pytesseract.image_to_string(image, config=config).strip()
                    if text:
                        all_results.append((name, text, len(text)))
                        self.logger.debug(f"{name}: extracted {len(text)} characters")
                except Exception as e:
                    self.logger.debug(f"{name} failed: {e}")
            
            # Return the longest text found
            if all_results:
                best_result = max(all_results, key=lambda x: x[2])
                self.logger.info(f"Best OCR result from method '{best_result[0]}': {best_result[2]} characters")
                return best_result[1]
            else:
                self.logger.info("No text extracted from image with any method")
                return ""
                
        except Exception as e:
            self.logger.error(f"All OCR attempts failed: {e}")
            return ""
    
    def process_docx(self, docx_path: str, diagnose: bool = True) -> Dict[str, any]:
        """Main method to process a DOCX file and extract all text content."""
        self.logger.info(f"Starting DOCX processing: {docx_path}")
        
        results = {
            'docx_path': docx_path,
            'native_text': {},
            'images': [],
            'image_ocr_text': [],
            'combined_text': '',
            'diagnosis': None
        }
        
        try:
            # Run diagnostics first if requested
            if diagnose:
                self.logger.info("Running DOCX diagnostics...")
                diagnosis = self.diagnose_docx(docx_path)
                results['diagnosis'] = diagnosis
            
            # Extract native text
            self.logger.info("Extracting native text...")
            native_text = self.extract_native_text(docx_path)
            results['native_text'] = native_text
            
            # Extract images and perform OCR
            self.logger.info("Extracting images and performing OCR...")
            images = self.extract_images_from_docx(docx_path)
            results['images'] = images
            
            # Process images with OCR
            image_ocr_results = []
            for img_data in images:
                try:
                    image = img_data["image"]
                    self.logger.info(f"Starting OCR for image {img_data['index']} (size: {image.size})")
                    
                    # Apply OCR
                    ocr_text = self.extract_text_from_image(image)
                    img_data["ocr_text"] = ocr_text
                    
                    image_ocr_results.append({
                        'index': img_data['index'],
                        'filename': img_data['filename'],
                        'size': img_data['size'],
                        'ocr_text': ocr_text
                    })
                    
                    if ocr_text.strip():
                        self.logger.info(f"✅ Extracted {len(ocr_text)} characters from image {img_data['index']}")
                    else:
                        self.logger.warning(f"❌ No text found in image {img_data['index']}")
                        
                except Exception as e:
                    self.logger.error(f"OCR failed for image {img_data['index']}: {e}")
                    img_data["ocr_text"] = ""
            
            results['image_ocr_text'] = image_ocr_results
            
            # Combine all extracted text
            combined_text = self._combine_extracted_text(results)
            results['combined_text'] = combined_text
            
            self.logger.info(f"DOCX processing completed. Total characters extracted: {len(combined_text)}")
            
        except Exception as e:
            self.logger.error(f"Error processing DOCX: {e}")
            raise
        
        return results
    
    def _combine_extracted_text(self, results: Dict[str, any]) -> str:
        """Combine all extracted text into a single string."""
        combined_parts = []
        
        native_text = results.get('native_text', {})
        
        # Add headers
        if native_text.get('headers'):
            combined_parts.append("=== HEADERS ===")
            combined_parts.extend(native_text['headers'])
            combined_parts.append("")
        
        # Add main paragraphs
        if native_text.get('paragraphs'):
            combined_parts.append("=== MAIN CONTENT ===")
            for para in native_text['paragraphs']:
                combined_parts.append(para['text'])
            combined_parts.append("")
        
        # Add tables
        if native_text.get('tables'):
            combined_parts.append("=== TABLES ===")
            for table in native_text['tables']:
                combined_parts.append(f"Table {table['table_num']} ({table['rows']}x{table['cols']}):")
                for row in table['data']:
                    combined_parts.append(" | ".join(row))
                combined_parts.append("")
        
        # Add image OCR text
        if results.get('image_ocr_text'):
            combined_parts.append("=== IMAGE OCR TEXT ===")
            for img_ocr in results['image_ocr_text']:
                if img_ocr['ocr_text'].strip():
                    combined_parts.append(f"[IMAGE {img_ocr['index']}] {img_ocr['filename']}")
                    combined_parts.append(img_ocr['ocr_text'])
                    combined_parts.append("")
        
        # Add footers
        if native_text.get('footers'):
            combined_parts.append("=== FOOTERS ===")
            combined_parts.extend(native_text['footers'])
        
        return '\n'.join(combined_parts)
    
    def save_extracted_text(self, results: Dict[str, any], output_filename: Optional[str] = None) -> str:
        """Save the extracted text to a file."""
        if not output_filename:
            docx_name = Path(results['docx_path']).stem
            output_filename = f"{docx_name}_extracted_text.txt"
        
        output_path = self.output_dir / output_filename
        
        # Create summary header
        import datetime
        summary = [
            f"DOCX TEXT EXTRACTION REPORT",
            f"Source DOCX: {results['docx_path']}",
            f"Total Paragraphs: {results['native_text'].get('total_paragraphs', 0)}",
            f"Total Tables: {results['native_text'].get('total_tables', 0)}",
            f"Images Found: {len(results['images'])}",
            f"Total Characters: {len(results['combined_text'])}",
            f"Extraction Date: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
            "="*80,
            ""
        ]
        
        # Write to file
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write('\n'.join(summary))
            f.write(results['combined_text'])
        
        self.logger.info(f"Extracted text saved to: {output_path}")
        return str(output_path)


# --- EasyOCR Extractor for Standalone Images ---
class EasyOCRExtractor:
    def __init__(self, enable_gpu: bool = True):
        try:
            import easyocr
            self.reader = easyocr.Reader(['en'], gpu=enable_gpu)
        except ImportError:
            raise ImportError("EasyOCR not installed. Install with: pip install easyocr")
    def preprocess_image(self, image_path: str) -> np.ndarray:
        image = cv2.imread(image_path)
        if image is None:
            raise ValueError(f"Could not load image: {image_path}")
        gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
        clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8, 8))
        enhanced = clahe.apply(gray)
        height, width = enhanced.shape[:2]
        if width < 500 or height < 500:
            scale = max(500 / width, 500 / height)
            enhanced = cv2.resize(enhanced, None, fx=scale, fy=scale, interpolation=cv2.INTER_CUBIC)
        return enhanced
    def extract_text(self, image_path: str) -> Dict:
        if not os.path.exists(image_path):
            raise FileNotFoundError(f"Image not found: {image_path}")
        try:
            results = self.reader.readtext(image_path)
            texts = []
            confidences = []
            for (bbox, text, confidence) in results:
                if confidence > 0.1:
                    texts.append(text.strip())
                    confidences.append(confidence)
            avg_confidence = sum(confidences) / len(confidences) if confidences else 0
            return {
                'text': '\n'.join(texts),
                'confidence': avg_confidence,
                'word_count': len(texts)
            }
        except Exception as e:
            return {'text': '', 'confidence': 0, 'error': str(e)}

# --- PDFTextExtractor class (PyMuPDF-based) ---
class PDFTextExtractor:
    """
    A robust PDF text extraction pipeline that combines native text extraction
    with OCR-based image text detection.
    """
    
    def __init__(self, tesseract_path: Optional[str] = None, output_dir: str = "extracted_texts"):
        """
        Initialize the PDF text extractor.
        
        Args:
            tesseract_path: Path to tesseract executable (if not in PATH)
            output_dir: Directory to save extracted text files
        """
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(exist_ok=True)
        
        # Configure tesseract path if provided
        if tesseract_path:
            pytesseract.pytesseract.tesseract_cmd = tesseract_path
        
        # Setup logging
        import logging
        logging.basicConfig(
            level=logging.INFO,
            format='%(asctime)s - %(levelname)s - %(message)s'
        )
        self.logger = logging.getLogger(__name__)
    
    def diagnose_pdf(self, pdf_path: str) -> Dict[str, any]:
        """
        Diagnose PDF structure to understand content types.
        
        Args:
            pdf_path: Path to the PDF file
            
        Returns:
            Dictionary with diagnostic information
        """
        try:
            import fitz  # PyMuPDF
            pdf_document = fitz.open(pdf_path)
            diagnosis = {
                'total_pages': len(pdf_document),
                'page_details': []
            }
            
            for page_num in range(len(pdf_document)):
                page = pdf_document[page_num]
                
                # Get basic page info
                page_info = {
                    'page_num': page_num + 1,
                    'mediabox': page.mediabox,
                    'rotation': page.rotation,
                    'image_count': len(page.get_images(full=True)),
                    'text_length': len(page.get_text()),
                    'drawings_count': 0,  # Initialize to 0
                    'links_count': len(page.get_links()),
                    'annotations_count': 0  # Initialize to 0
                }
                
                # Safely get drawings count
                try:
                    drawings = page.get_drawings()
                    page_info['drawings_count'] = len(list(drawings))  # Convert generator to list
                except:
                    page_info['drawings_count'] = 0
                
                # Safely get annotations count
                try:
                    annotations = page.annots()
                    page_info['annotations_count'] = len(list(annotations))  # Convert generator to list
                except:
                    page_info['annotations_count'] = 0
                
                # Check for different content types
                images = page.get_images(full=True)
                page_info['image_details'] = []
                
                for img_idx, img in enumerate(images):
                    try:
                        xref = img[0]
                        base_image = pdf_document.extract_image(xref)
                        img_detail = {
                            'index': img_idx + 1,
                            'xref': xref,
                            'ext': base_image.get('ext', 'unknown'),
                            'width': base_image.get('width', 0),
                            'height': base_image.get('height', 0),
                            'colorspace': base_image.get('colorspace', 'unknown'),
                            'size_bytes': len(base_image.get('image', b''))
                        }
                        page_info['image_details'].append(img_detail)
                    except Exception as e:
                        page_info['image_details'].append({
                            'index': img_idx + 1,
                            'error': str(e)
                        })
                
                diagnosis['page_details'].append(page_info)
                
                self.logger.info(f"Page {page_num + 1}: {page_info['text_length']} chars text, "
                               f"{page_info['image_count']} images, {page_info['drawings_count']} drawings")
            
            pdf_document.close()
            return diagnosis
            
        except Exception as e:
            self.logger.error(f"PDF diagnosis failed: {e}")
            return {'error': str(e)}
    
    def extract_native_text(self, pdf_document: fitz.Document) -> List[Dict]:
        """
        Extract native text from PDF pages using PyMuPDF with positional information.
        
        Args:
            pdf_document: PyMuPDF document object
            
        Returns:
            List of dictionaries containing text content and position for each page
        """
        page_texts = []
        
        for page_num in range(len(pdf_document)):
            page = pdf_document[page_num]
            
            # Extract text with position information
            text_dict = page.get_text("dict")
            
            # Also get simple text for fallback
            simple_text = page.get_text("text")
            
            # Process text blocks to get position information
            text_blocks = []
            
            for block in text_dict.get("blocks", []):
                if "lines" in block:  # Text block
                    block_text = ""
                    for line in block["lines"]:
                        for span in line["spans"]:
                            block_text += span["text"]
                        block_text += "\n"
                    
                    if block_text.strip():
                        text_blocks.append({
                            "text": block_text.strip(),
                            "bbox": block["bbox"],  # (x0, y0, x1, y1)
                            "type": "text"
                        })
            
            page_data = {
                "page_num": page_num,
                "simple_text": simple_text,
                "text_blocks": text_blocks,
                "page_height": page.rect.height,
                "page_width": page.rect.width
            }
            
            page_texts.append(page_data)
            self.logger.info(f"Extracted native text from page {page_num + 1} with {len(text_blocks)} text blocks")
        
        return page_texts
    
    def extract_images_from_pdf(self, pdf_document: fitz.Document) -> List[Tuple[int, List[Dict]]]:
        """
        Enhanced image extraction from PDF pages with position information.
        
        Args:
            pdf_document: PyMuPDF document object
            
        Returns:
            List of tuples containing (page_number, list_of_image_dicts_with_position)
        """
        all_images = []
        
        # Create extracted_images folder
        extracted_images_dir = self.output_dir / "extracted_images"
        extracted_images_dir.mkdir(exist_ok=True)
        
        for page_num in range(len(pdf_document)):
            page = pdf_document[page_num]
            page_images = []
            
            # Method 1: Extract images using get_images()
            image_list = page.get_images(full=True)
            self.logger.info(f"Page {page_num + 1}: Found {len(image_list)} images using get_images()")
            
            for img_index, img in enumerate(image_list):
                try:
                    # Get image data
                    xref = img[0]
                    base_image = pdf_document.extract_image(xref)
                    image_bytes = base_image["image"]
                    
                    # Filter out tiny images (likely decorative elements)
                    if len(image_bytes) < 1000:  # Skip very small images
                        self.logger.debug(f"Skipping tiny image {img_index + 1} on page {page_num + 1} (size: {len(image_bytes)} bytes)")
                        continue
                    
                    # Convert to PIL Image
                    image = Image.open(io.BytesIO(image_bytes))
                    
                    # Filter out very small images by dimensions
                    if image.width < 50 or image.height < 50:
                        self.logger.debug(f"Skipping small image {img_index + 1} on page {page_num + 1} (dimensions: {image.width}x{image.height})")
                        continue
                    
                    # Save the image to disk
                    img_filename = f"page{page_num + 1}_img{img_index + 1}.png"
                    img_path = extracted_images_dir / img_filename
                    image.save(img_path)
                    
                    # Get image position on the page
                    img_bbox = None
                    try:
                        # Find image position using transformation matrix
                        for item in page.get_images(full=True):
                            if item[0] == xref:
                                # Get image rect
                                img_rects = page.get_image_rects(item)
                                if img_rects:
                                    img_bbox = img_rects[0]  # Use first rect if multiple
                                break
                    except:
                        # Fallback: estimate position (top of page)
                        img_bbox = fitz.Rect(0, 0, image.width, image.height)
                    
                    image_data = {
                        "image": image,
                        "bbox": img_bbox,
                        "index": img_index + 1,
                        "size": (image.width, image.height),
                        "type": "image",
                        "saved_path": str(img_path),
                        "filename": img_filename
                    }
                    
                    page_images.append(image_data)
                    self.logger.info(f"Successfully extracted and saved image {img_index + 1} from page {page_num + 1} (size: {image.width}x{image.height}, {len(image_bytes)} bytes) -> {img_filename}")
                    
                except Exception as e:
                    self.logger.warning(f"Failed to extract image {img_index + 1} from page {page_num + 1}: {e}")
            
            all_images.append((page_num, page_images))
            self.logger.info(f"Page {page_num + 1}: Successfully extracted {len(page_images)} valid images")
        
        return all_images
    
    def preprocess_image_for_ocr(self, image: Image.Image) -> Image.Image:
        """
        Enhanced image preprocessing to improve OCR accuracy.
        
        Args:
            image: PIL Image object
            
        Returns:
            Preprocessed PIL Image
        """
        try:
            # Convert to RGB if necessary
            if image.mode != 'RGB':
                image = image.convert('RGB')
            
            # Convert PIL to OpenCV format
            opencv_image = cv2.cvtColor(np.array(image), cv2.COLOR_RGB2BGR)
            
            # Convert to grayscale
            gray = cv2.cvtColor(opencv_image, cv2.COLOR_BGR2GRAY)
            
            # Apply multiple preprocessing techniques and return the best one
            processed_images = []
            
            # Method 1: Basic preprocessing
            denoised = cv2.medianBlur(gray, 3)
            thresh1 = cv2.adaptiveThreshold(
                denoised, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, 11, 2
            )
            processed_images.append(('adaptive_thresh', Image.fromarray(thresh1)))
            
            # Method 2: OTSU thresholding
            _, thresh2 = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
            processed_images.append(('otsu_thresh', Image.fromarray(thresh2)))
            
            # Method 3: Simple thresholding
            _, thresh3 = cv2.threshold(gray, 127, 255, cv2.THRESH_BINARY)
            processed_images.append(('simple_thresh', Image.fromarray(thresh3)))
            
            # Method 4: Enhanced contrast
            clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8,8))
            enhanced = clahe.apply(gray)
            _, thresh4 = cv2.threshold(enhanced, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
            processed_images.append(('enhanced_contrast', Image.fromarray(thresh4)))
            
            # Method 5: Morphological operations
            kernel = np.ones((2,2), np.uint8)
            morph = cv2.morphologyEx(gray, cv2.MORPH_CLOSE, kernel)
            _, thresh5 = cv2.threshold(morph, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
            processed_images.append(('morphological', Image.fromarray(thresh5)))
            
            # Return the adaptive threshold as default (usually works well)
            return processed_images[0][1]
            
        except Exception as e:
            self.logger.warning(f"Image preprocessing failed: {e}, using original image")
            return image
    
    def extract_text_from_image(self, image: Image.Image) -> str:
        """
        Extract text from image using Tesseract OCR with multiple approaches.
        
        Args:
            image: PIL Image object
            
        Returns:
            Extracted text string
        """
        try:
            # Test if Tesseract is available
            try:
                version = pytesseract.get_tesseract_version()
                self.logger.debug(f"Using Tesseract version: {version}")
            except Exception as e:
                self.logger.error(f"Tesseract not available: {e}")
                return ""
            
            all_results = []
            
            # Try 1: Original image with different PSM modes
            original_configs = [
                ('original_psm6', r'--oem 3 --psm 6'),
                ('original_psm7', r'--oem 3 --psm 7'),
                ('original_psm8', r'--oem 3 --psm 8'),
                ('original_psm11', r'--oem 3 --psm 11'),
                ('original_psm13', r'--oem 3 --psm 13')
            ]
            
            for name, config in original_configs:
                try:
                    text = pytesseract.image_to_string(image, config=config).strip()
                    if text:
                        all_results.append((name, text, len(text)))
                        self.logger.debug(f"{name}: extracted {len(text)} characters")
                except Exception as e:
                    self.logger.debug(f"{name} failed: {e}")
            
            # Try 2: Preprocessed images
            try:
                # Convert to RGB if necessary
                if image.mode != 'RGB':
                    rgb_image = image.convert('RGB')
                else:
                    rgb_image = image
                
                # Convert to numpy array
                img_array = np.array(rgb_image)
                
                # Try different preprocessing approaches
                preprocessing_methods = [
                    ('grayscale', self._to_grayscale),
                    ('threshold_otsu', self._apply_otsu_threshold),
                    ('threshold_adaptive', self._apply_adaptive_threshold),
                    ('enhance_contrast', self._enhance_contrast),
                    ('remove_noise', self._remove_noise),
                    ('resize_2x', self._resize_image),
                ]
                
                for method_name, method_func in preprocessing_methods:
                    try:
                        processed_img = method_func(img_array)
                        processed_pil = Image.fromarray(processed_img)
                        
                        # Try OCR with different configs
                        configs = [r'--oem 3 --psm 6', r'--oem 3 --psm 7', r'--oem 3 --psm 11']
                        for config in configs:
                            try:
                                text = pytesseract.image_to_string(processed_pil, config=config).strip()
                                if text:
                                    all_results.append((f"{method_name}_{config.split()[-1]}", text, len(text)))
                                    self.logger.debug(f"{method_name} with {config}: extracted {len(text)} characters")
                            except:
                                continue
                                
                    except Exception as e:
                        self.logger.debug(f"Preprocessing method {method_name} failed: {e}")
                        continue
            
            except Exception as e:
                self.logger.warning(f"Preprocessing attempts failed: {e}")
            
            # Try 3: Scale the image up (sometimes helps with small text)
            try:
                # Scale up by 2x
                width, height = image.size
                scaled_image = image.resize((width * 2, height * 2), Image.LANCZOS)
                
                configs = [r'--oem 3 --psm 6', r'--oem 3 --psm 7']
                for config in configs:
                    try:
                        text = pytesseract.image_to_string(scaled_image, config=config).strip()
                        if text:
                            all_results.append((f"scaled_{config.split()[-1]}", text, len(text)))
                            self.logger.debug(f"Scaled image with {config}: extracted {len(text)} characters")
                    except:
                        continue
                        
            except Exception as e:
                self.logger.debug(f"Image scaling failed: {e}")
            
            # Return the longest text found
            if all_results:
                best_result = max(all_results, key=lambda x: x[2])
                self.logger.info(f"Best OCR result from method '{best_result[0]}': {best_result[2]} characters")
                return best_result[1]
            else:
                self.logger.info("No text extracted from image with any method")
                return ""
                
        except Exception as e:
            self.logger.error(f"All OCR attempts failed: {e}")
            return ""
    
    def _to_grayscale(self, img_array):
        """Convert image to grayscale."""
        if len(img_array.shape) == 3:
            return cv2.cvtColor(img_array, cv2.COLOR_RGB2GRAY)
        return img_array
    
    def _apply_otsu_threshold(self, img_array):
        """Apply OTSU thresholding."""
        gray = self._to_grayscale(img_array)
        _, thresh = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
        return thresh
    
    def _apply_adaptive_threshold(self, img_array):
        """Apply adaptive thresholding."""
        gray = self._to_grayscale(img_array)
        return cv2.adaptiveThreshold(gray, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, 11, 2)
    
    def _enhance_contrast(self, img_array):
        """Enhance image contrast."""
        gray = self._to_grayscale(img_array)
        clahe = cv2.createCLAHE(clipLimit=3.0, tileGridSize=(8,8))
        enhanced = clahe.apply(gray)
        _, thresh = cv2.threshold(enhanced, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
        return thresh
    
    def _remove_noise(self, img_array):
        """Remove noise from image."""
        gray = self._to_grayscale(img_array)
        denoised = cv2.medianBlur(gray, 3)
        _, thresh = cv2.threshold(denoised, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
        return thresh
    
    def _resize_image(self, img_array):
        """Resize image by 2x."""
        height, width = img_array.shape[:2]
        resized = cv2.resize(img_array, (width * 2, height * 2), interpolation=cv2.INTER_CUBIC)
        return self._to_grayscale(resized)
    
    def combine_text_and_images_by_position(self, page_text_data: Dict, page_images: List[Dict]) -> List[Dict]:
        """
        Combine text blocks and images based on their position on the page to maintain reading order.
        
        Args:
            page_text_data: Dictionary containing text blocks and page info
            page_images: List of image dictionaries with position info
            
        Returns:
            List of content blocks sorted by reading order (top to bottom, left to right)
        """
        content_blocks = []
        
        # Add text blocks
        for text_block in page_text_data.get("text_blocks", []):
            content_blocks.append({
                "type": "text",
                "content": text_block["text"],
                "bbox": text_block["bbox"],
                "y_position": text_block["bbox"][1]  # y0 coordinate for sorting
            })
        
        # Add image blocks (with OCR text if available)
        for img_data in page_images:
            bbox = img_data.get("bbox")
            if bbox:
                y_pos = bbox.y0 if hasattr(bbox, 'y0') else bbox[1] if isinstance(bbox, (list, tuple)) else 0
            else:
                y_pos = 0  # Default to top if no position info
            
            content_blocks.append({
                "type": "image",
                "content": img_data,
                "bbox": bbox,
                "y_position": y_pos
            })
        
        # Sort by vertical position (top to bottom)
        content_blocks.sort(key=lambda x: x["y_position"])
        
        return content_blocks
    
    def extract_text_from_pdf_as_images(self, pdf_path: str) -> List[str]:
        """
        Convert PDF pages to images and extract text using OCR.
        This is useful for scanned PDFs or when native text extraction fails.
        
        Args:
            pdf_path: Path to the PDF file
            
        Returns:
            List of text content for each page
        """
        try:
            from pdf2image import convert_from_path
            
            # Convert PDF pages to images
            images = convert_from_path(pdf_path, dpi=300)
            page_texts = []
            
            for i, image in enumerate(images):
                text = self.extract_text_from_image(image)
                page_texts.append(text)
                self.logger.info(f"Extracted OCR text from PDF page {i + 1} (as image)")
            
            return page_texts
            
        except Exception as e:
            self.logger.error(f"Failed to convert PDF to images: {e}")
            return []
    
    def process_pdf(self, pdf_path: str, use_fallback_ocr: bool = True, diagnose: bool = True) -> Dict[str, any]:
        """
        Main method to process a PDF file and extract all text content.
        
        Args:
            pdf_path: Path to the PDF file
            use_fallback_ocr: Whether to use full-page OCR as fallback
            diagnose: Whether to run PDF diagnostics first
            
        Returns:
            Dictionary containing extraction results
        """
        self.logger.info(f"Starting PDF processing: {pdf_path}")
        
        results = {
            'pdf_path': pdf_path,
            'native_text': [],
            'image_ocr_text': [],
            'fallback_ocr_text': [],
            'combined_text': '',
            'total_pages': 0,
            'images_found': 0,
            'diagnosis': None
        }
        
        try:
            # Run diagnostics first if requested
            if diagnose:
                self.logger.info("Running PDF diagnostics...")
                diagnosis = self.diagnose_pdf(pdf_path)
                results['diagnosis'] = diagnosis
                
                # Log summary
                total_images = sum(page['image_count'] for page in diagnosis.get('page_details', []))
                self.logger.info(f"Diagnosis complete: {diagnosis.get('total_pages', 0)} pages, {total_images} total images detected")
            
            # Open PDF document
            pdf_document = fitz.open(pdf_path)
            results['total_pages'] = len(pdf_document)
            
            # Extract native text with position information
            self.logger.info("Extracting native text...")
            native_texts = self.extract_native_text(pdf_document)
            results['native_text'] = native_texts
            
            # Extract images and perform OCR with position information
            self.logger.info("Extracting images and performing OCR...")
            page_images = self.extract_images_from_pdf(pdf_document)
            
            # Process each page to maintain reading order
            ordered_content = []
            total_images = 0
            
            for page_num, images in page_images:
                page_image_count = len(images)
                total_images += page_image_count
                
                self.logger.info(f"Processing {page_image_count} images from page {page_num + 1}")
                
                # Process images and extract OCR text
                processed_images = []
                for img_idx, image_data in enumerate(images):
                    try:
                        image = image_data["image"]
                        self.logger.info(f"Starting OCR for image {img_idx + 1} on page {page_num + 1} (size: {image.size})")
                        
                        # Try enhanced OCR without debug output
                        ocr_text = self.extract_text_from_image(image)
                        
                        # Add OCR text to image data
                        image_data["ocr_text"] = ocr_text
                        processed_images.append(image_data)
                        
                        if ocr_text.strip():
                            self.logger.info(f"✅ Extracted {len(ocr_text)} characters from image {img_idx + 1} on page {page_num + 1}")
                        else:
                            self.logger.warning(f"❌ No text found in image {img_idx + 1} on page {page_num + 1}")
                            
                            # For debugging, save problematic image only if needed
                            if page_num == 0 and img_idx == 0:  # Only for first image to avoid spam
                                self.logger.debug("First image had no OCR results - this may indicate OCR configuration issues")
                                    
                    except Exception as e:
                        self.logger.error(f"OCR failed for image {img_idx + 1} on page {page_num + 1}: {e}")
                        image_data["ocr_text"] = ""
                        processed_images.append(image_data)
                
                # Combine text and images by position for this page
                if page_num < len(native_texts):
                    page_content = self.combine_text_and_images_by_position(
                        native_texts[page_num], processed_images
                    )
                    ordered_content.append(page_content)
                else:
                    # Page with only images
                    ordered_content.append([{
                        "type": "image",
                        "content": img_data,
                        "bbox": img_data.get("bbox"),
                        "y_position": 0
                    } for img_data in processed_images])
            
            results['ordered_content'] = ordered_content
            results['images_found'] = total_images
            
            # Close the PDF document
            pdf_document.close()
            
            # Fallback: Use full-page OCR if native text extraction yields poor results
            if use_fallback_ocr:
                total_native_chars = 0
                for page_data in native_texts:
                    if isinstance(page_data, dict):
                        total_native_chars += len(page_data.get('simple_text', ''))
                    else:
                        total_native_chars += len(str(page_data))
                
                if total_native_chars < 100:  # Threshold for "poor" native extraction
                    self.logger.info("Native text extraction yielded minimal results. Using full-page OCR...")
                    fallback_texts = self.extract_text_from_pdf_as_images(pdf_path)
                    results['fallback_ocr_text'] = fallback_texts
            
            # Combine all extracted text
            combined_text = self._combine_extracted_text(results)
            results['combined_text'] = combined_text
            
            self.logger.info(f"PDF processing completed. Total characters extracted: {len(combined_text)}")
            
        except Exception as e:
            self.logger.error(f"Error processing PDF: {e}")
            raise
        
        return results
    
    def _combine_extracted_text(self, results: Dict[str, any]) -> str:
        """
        Combine all extracted text into a single string, maintaining reading order.
        
        Args:
            results: Dictionary containing extraction results
            
        Returns:
            Combined text string with preserved reading order
        """
        combined_parts = []
        
        # Use ordered content if available (new method)
        if 'ordered_content' in results and results['ordered_content']:
            for page_num, page_content in enumerate(results['ordered_content']):
                if not page_content:
                    continue
                    
                combined_parts.append(f"\n=== PAGE {page_num + 1} ===")
                
                for block in page_content:
                    if block["type"] == "text":
                        combined_parts.append(block["content"])
                    elif block["type"] == "image":
                        img_data = block["content"]
                        ocr_text = img_data.get("ocr_text", "")
                        if ocr_text.strip():
                            combined_parts.append(f"\n[IMAGE {img_data.get('index', '?')}]")
                            combined_parts.append(ocr_text)
                        else:
                            combined_parts.append(f"\n[IMAGE {img_data.get('index', '?')} - No text detected]")
                
                combined_parts.append("\n" + "="*50 + "\n")
        
        # Fallback to old method if ordered content not available
        else:
            for page_num in range(results['total_pages']):
                page_parts = []
                
                # Add native text
                if page_num < len(results.get('native_text', [])):
                    if isinstance(results['native_text'][page_num], dict):
                        # New format with position info
                        simple_text = results['native_text'][page_num].get('simple_text', '')
                        if simple_text.strip():
                            page_parts.append("=== NATIVE TEXT ===")
                            page_parts.append(simple_text)
                    elif isinstance(results['native_text'][page_num], str):
                        # Old format
                        if results['native_text'][page_num].strip():
                            page_parts.append("=== NATIVE TEXT ===")
                            page_parts.append(results['native_text'][page_num])
                
                # Add image OCR text
                if (results.get('image_ocr_text') and 
                    page_num < len(results['image_ocr_text']) and 
                    results['image_ocr_text'][page_num].strip()):
                    page_parts.append("=== IMAGE OCR TEXT ===")
                    page_parts.append(results['image_ocr_text'][page_num])
                
                # Add fallback OCR text
                if (results.get('fallback_ocr_text') and 
                    page_num < len(results['fallback_ocr_text']) and 
                    results['fallback_ocr_text'][page_num].strip()):
                    page_parts.append("=== FALLBACK OCR TEXT ===")
                    page_parts.append(results['fallback_ocr_text'][page_num])
                
                if page_parts:
                    combined_parts.append(f"\n=== PAGE {page_num + 1} ===")
                    combined_parts.extend(page_parts)
                    combined_parts.append("\n" + "="*50 + "\n")
        
        return '\n'.join(combined_parts)
    
    def save_extracted_text(self, results: Dict[str, any], output_filename: Optional[str] = None) -> str:
        """
        Save the extracted text to a file.
        
        Args:
            results: Dictionary containing extraction results
            output_filename: Custom output filename (optional)
            
        Returns:
            Path to the saved file
        """
        if not output_filename:
            pdf_name = Path(results['pdf_path']).stem
            output_filename = f"{pdf_name}_extracted_text.txt"
        
        output_path = self.output_dir / output_filename
        
        # Create summary header
        import logging
        summary = [
            f"PDF TEXT EXTRACTION REPORT",
            f"Source PDF: {results['pdf_path']}",
            f"Total Pages: {results['total_pages']}",
            f"Images Found: {results['images_found']}",
            f"Total Characters: {len(results['combined_text'])}",
            f"Extraction Date: {logging.Formatter().formatTime(logging.LogRecord('', 0, '', 0, '', (), None))}",
            "="*80,
            ""
        ]
        
        # Write to file
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write('\n'.join(summary))
            f.write(results['combined_text'])
        
        self.logger.info(f"Extracted text saved to: {output_path}")
        return str(output_path)


# --- Main function to handle both PDFs and images ---
def main():
    # Set your Tesseract path if needed
    tesseract_path = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
    file_path = "claims/67a0d2d07d36232f470e0ee2/22JAN VS479 fuelling issues YES_1740483912.docx"  # Change this to your file (PDF or image)
    ext = os.path.splitext(file_path)[1].lower()
    if ext == '.pdf':
        print(f"\n[PDF Mode] Processing: {file_path}")
        extractor = PDFTextExtractor(tesseract_path=tesseract_path)
        results = extractor.process_pdf(file_path)
        print("\n=== PDF Native Text ===")
        for page in results['native_text']:
            print(f"\n--- Page {page['page_num']} ---\n{page['text']}")
        print("\n=== PDF Image OCR Results ===")
        if 'ordered_content' in results:
            for page_num, page_content in enumerate(results['ordered_content']):
                for block in page_content:
                    if block['type'] == 'image':
                        img_data = block['content']
                        filename = img_data.get('filename', 'unknown')
                        ocr_text = img_data.get('ocr_text', '')
                        print(f"Page {page_num + 1} Image {img_data['index']} (saved as: {filename}):")
                        print(f"Size: {img_data['size']}")
                        if ocr_text:
                            print(f"OCR Text: {ocr_text[:100]}{'...' if len(ocr_text) > 100 else ''}")
                        else:
                            print("OCR Text: No text detected")
                        print()
        else:
            print("No image extraction results available")
    elif ext == '.docx':
        print(f"\n[DOCX Mode] Processing: {file_path}")
        extractor = DocxTextExtractor(tesseract_path=tesseract_path)
        results = extractor.process_docx(file_path)
        
        # Display summary
        native_text = results['native_text']
        print(f"\nDOCX Summary:")
        print(f"  Paragraphs: {native_text.get('total_paragraphs', 0)}")
        print(f"  Tables: {native_text.get('total_tables', 0)}")
        print(f"  Images: {len(results['images'])}")
        print(f"  Headers: {len(native_text.get('headers', []))}")
        print(f"  Footers: {len(native_text.get('footers', []))}")
        
        # Display first few paragraphs
        print("\n=== DOCX Native Text (First 3 Paragraphs) ===")
        for i, para in enumerate(native_text.get('paragraphs', [])[:3]):
            print(f"Paragraph {para['paragraph_num']}: {para['text'][:200]}{'...' if len(para['text']) > 200 else ''}")
        
        # Display tables
        if native_text.get('tables'):
            print(f"\n=== DOCX Tables ({len(native_text['tables'])}) ===")
            for table in native_text['tables'][:2]:  # Show first 2 tables
                print(f"Table {table['table_num']} ({table['rows']}x{table['cols']}):")
                for row in table['data'][:3]:  # Show first 3 rows
                    print("  " + " | ".join(row))
                if table['rows'] > 3:
                    print("  ...")
        
        # Display image OCR results
        print("\n=== DOCX Image OCR Results ===")
        for img_ocr in results['image_ocr_text']:
            filename = img_ocr['filename']
            ocr_text = img_ocr['ocr_text']
            print(f"Image {img_ocr['index']} (saved as: {filename}):")
            print(f"Size: {img_ocr['size']}")
            if ocr_text:
                print(f"OCR Text: {ocr_text[:100]}{'...' if len(ocr_text) > 100 else ''}")
            else:
                print("OCR Text: No text detected")
            print()
        
        # Save extracted text
        extractor.save_extracted_text(results)
        
    elif ext in ['.png', '.jpg', '.jpeg', '.tiff', '.bmp']:
        print(f"\n[Image Mode] Processing: {file_path}")
        extractor = EasyOCRExtractor(enable_gpu=False)
        result = extractor.extract_text(file_path)
        print(f"Confidence: {result['confidence']:.2f}")
        print(f"Words extracted: {result.get('word_count', 0)}")
        print("\nExtracted Text:")
        print(result['text'])
        if 'error' in result:
            print(f"Error: {result['error']}")
        
        # Save extracted text to file
        if result['text'].strip():
            output_dir = Path("extracted_texts")
            output_dir.mkdir(exist_ok=True)
            
            # Create filename based on input image
            image_name = Path(file_path).stem
            output_filename = f"{image_name}_extracted_text.txt"
            output_path = output_dir / output_filename
            
            # Create summary header
            import datetime
            summary = [
                f"IMAGE TEXT EXTRACTION REPORT",
                f"Source Image: {file_path}",
                f"Confidence: {result['confidence']:.2f}",
                f"Words Extracted: {result.get('word_count', 0)}",
                f"Total Characters: {len(result['text'])}",
                f"Extraction Date: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
                "="*80,
                "",
                result['text']
            ]
            
            # Write to file
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write('\n'.join(summary))
            
            print(f"\nExtracted text saved to: {output_path}")
        else:
            print("\nNo text extracted - nothing to save")
    else:
        print(f"Unsupported file type: {ext}")

if __name__ == "__main__":
    main()